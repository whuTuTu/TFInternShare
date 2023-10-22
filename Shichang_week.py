# -*- coding: utf-8 -*-
"""
生成市场周报的文字部分
红色文字指的是需要替换的
"""
# 载入包
from iFinDPy import *
from threading import Thread, Lock, Semaphore
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
import datetime
import math
import add_float_picture
import warnings
warnings.filterwarnings(action='ignore')  # 导入warnings模块，并指定忽略代码运行中的警告信息

# -------------------------------------------- 修改此处日期 --------------------------------------------
date = datetime.date.today()  # 在周六或者周日做周报，直接运行，会自动找到本周五的日期
# date = date(2023, 10, 14)  # 指定特定周在此处修改日期可以填写周六或者周日的日期

# -------------------------------------------- 以下是代码 --------------------------------------------
nearest_friday = date + datetime.timedelta(4 - date.weekday())
print(nearest_friday)
nearest_monday = date + datetime.timedelta(0 - date.weekday())
print(nearest_monday)
format = "%m.%d"
format1 = "20%y.%m.%d"
format2 = "20%y-%m-%d"
format3 = "2023-01-01,20%y-%m-%d"
format4 = "20%y-%m-%d,2013-%m-%d,20%y-%m-%d,101,100"
Friday = nearest_friday.strftime(format1)
Monday = nearest_monday.strftime(format1)
friday = nearest_friday.strftime(format)
monday = nearest_monday.strftime(format)
Mon = nearest_monday.strftime(format2)
Fri = nearest_friday.strftime(format2)
Year = nearest_friday.strftime(format3)
Value = nearest_friday.strftime(format4)
sem = Semaphore(5)
dllock = Lock()
thsLogin = THS_iFinDLogin("tfzq1556", "752862")
document = Document()
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)


def head1(str):
    paragrapha = document.add_heading(str, level=1)
    paragrapha.style.font.size = Pt(16)
    paragrapha.style.font.color.rgb = RGBColor(0, 0, 0)
    paragrapha.paragraph_format.space_after = Pt(13)
    paragrapha.paragraph_format.space_before = Pt(13)


def head2(str):
    paragrapha = document.add_heading(str, level=2)
    paragrapha.style.font.size = Pt(14)
    paragrapha.style.font.color.rgb = RGBColor(0, 0, 0)
    paragrapha.paragraph_format.space_after = Pt(13)
    paragrapha.paragraph_format.space_before = Pt(13)


def head3(str):
    paragrapha = document.add_paragraph(str, style='Body Text 3')
    paragrapha.style.font.size = Pt(12)
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrapha.paragraph_format.space_after = Pt(4)
    paragrapha.paragraph_format.space_before = Pt(20)
    paragrapha.paragraph_format.line_spacing = 1.08
    for run in paragrapha.runs:
        run.bold = True
    return paragrapha


def head4(str):
    paragrapha = document.add_paragraph(str, style='Body Text 3')
    paragrapha.style.font.size = Pt(12)
    paragrapha.paragraph_format.first_line_indent = Cm(0.8)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.paragraph_format.line_spacing = 1.5
    for run in paragrapha.runs:
        run.bold = True
    return paragrapha


def note():
    date = datetime.date.today()
    nearest_friday = date + datetime.timedelta(4 - date.weekday())
    format = "20%y.%m.%d"
    friday = nearest_friday.strftime(format)
    paragrapha = document.add_paragraph('数据来源：Wind，金融产品研发部整理，数据截止到', style='Body Text 2')
    paragrapha.add_run(str(friday))
    paragrapha.style.font.size = Pt(10)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.paragraph_format.line_spacing = 1.5
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragrapha.runs:
        run.bold = True
    return paragrapha


def note1():
    date = datetime.date.today()
    nearest_friday = date + datetime.timedelta(4 - date.weekday())
    format = "20%y.%m.%d"
    friday = nearest_friday.strftime(format)
    paragrapha = document.add_paragraph('数据来源：朝阳永续，金融产品研发部整理，数据截止到', style='Body Text 2')
    paragrapha.add_run(str(friday))
    paragrapha.style.font.size = Pt(10)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.paragraph_format.line_spacing = 1.5
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragrapha.runs:
        run.bold = True
    return paragrapha


def note2():
    date = datetime.date.today()
    nearest_friday = date + datetime.timedelta(4 - date.weekday())
    format = "20%y.%m.%d"
    friday = nearest_friday.strftime(format)
    paragrapha = document.add_paragraph('所有数据来源：Wind，金融产品研发部整理，数据截止到', style='Body Text 2')
    paragrapha.add_run(str(friday))
    paragrapha.style.font.size = Pt(10)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.paragraph_format.line_spacing = 1.5
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragrapha.runs:
        run.bold = True
    return paragrapha


def text(str):
    paragrapha = document.add_paragraph(str)
    paragrapha.style.font.size = Pt(12)
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrapha.paragraph_format.first_line_indent = Cm(0.8)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.paragraph_format.line_spacing = 1.5
    return paragrapha


def photo():
    paragrapha = document.add_paragraph('图片', style='Body Text')
    paragrapha.style.font.size = Pt(12)
    paragrapha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrapha.paragraph_format.space_after = Pt(4)
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.line_spacing = 1.08
    paragrapha.style.font.color.rgb = RGBColor(255, 0, 0)


p = document.add_paragraph()
add_float_picture.add_float_picture(p, r'input/市场周报.jpg', width=Cm(21))
p.add_run(str(Friday))

document.add_section(WD_SECTION.NEW_PAGE)
section = document.sections[0]
section.top_margin = Cm(18.5)
section.left_margin = Cm(1.55)
header_para = section.header.paragraphs[-1]
add_float_picture.add_float_picture(header_para, r'input/back.jpg', width=Cm(21))

#################################宽基指数###############################################
head1('一、市场表现周度回顾')
head2('（一）股票市场')
photo()
note()
up = text('本周')
up.add_run('（').bold = True
up.add_run(str(monday)).bold = True
up.add_run('-').bold = True
up.add_run(str(friday)).bold = True
up.add_run('）').bold = True
up.add_run('各大宽基指数')
index = THS_HQ('000001.SH,399001.SZ,399006.SZ,000016.SH,000300.SH,000905.SH,000852.SH', 'changeRatio', 'Interval:W',
               Mon, Fri).data
index['thscode'] = ['上证指数', '深证成指', '创业板指', '上证50', '沪深300', '中证500', '中证1000']
print(index)
B = index['changeRatio']
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
index.sort_values(by='changeRatio', ascending=False, inplace=True)
first_name = index.iloc[0, :]['thscode']
first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
last_name = index.iloc[-1, :]['thscode']
last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'])
if n2 == 1:  # 只有一个跌
    # 降序,筛选极值
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    up.add_run('多数上涨。其中，')
    up.add_run(str(first_name))
    up.add_run('领涨，周涨')
    up.add_run(str(first_score))
    up.add_run('；仅')
    up.add_run(str(last_name))
    up.add_run('下跌，周跌')
    up.add_run(str(last_score))
elif n1 == 1:  # 只有一个涨
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    up.add_run('多数下跌。其中，')
    up.add_run(str(first_name))
    up.add_run('领跌，周跌')
    up.add_run(str(first_score))
    up.add_run('；仅')
    up.add_run(str(last_name))
    up.add_run('上涨，周涨')
    up.add_run(str(last_score))
elif n2 == 0:  # 没有跌
    up.add_run('全部上涨。其中，')
    up.add_run(str(first_name))
    up.add_run('领涨，周涨')
    up.add_run(str(first_score))
    up.add_run('；')
    up.add_run(str(last_name))
    up.add_run('涨幅最小，周涨')
    up.add_run(str(last_score))
elif n1 == 0:  # 没有涨
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'] * -1)
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    up.add_run('全部下跌。其中，')
    up.add_run(str(first_name))
    up.add_run('领跌，周跌')
    up.add_run(str(first_score))
    up.add_run('；')
    up.add_run(str(last_name))
    up.add_run('跌幅最小，周跌')
    up.add_run(str(last_score))
else:
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    up.add_run('涨跌不一。其中，')
    up.add_run(str(first_name))
    up.add_run('领涨，周涨')
    up.add_run(str(first_score))
    up.add_run('；')
    up.add_run(str(last_name))
    up.add_run('领跌，周跌')
    up.add_run(str(last_score))
index = THS_BD('000001.SH,399001.SZ,399006.SZ,000016.SH,000300.SH,000905.SH,000852.SH', 'ths_int_chg_ratio_index',
               Year).data
index['thscode'] = ['上证指数', '深证成指', '创业板指', '上证50', '沪深300', '中证500', '中证1000']
up.add_run('。年度来看，今年以来各大指数')
B = index['ths_int_chg_ratio_index']
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
if n2 == 0:
    up.add_run('全部呈上涨状态。')
elif n1 == 0:
    up.add_run('全部呈下跌状态。')
elif n1 > n2:
    up.add_run('多数呈上涨状态。')
elif n2 > n1:
    up.add_run('多数呈下跌状态。')
else:
    up.add_run('有涨有跌。')
up.add_run('指数估值方面，')
index = THS_BD('000001.SH,399001.SZ,399006.SZ,000016.SH,000300.SH,000905.SH,000852.SH', 'ths_pe_ttm_quantile_index',
               Value).data
index['thscode'] = ['上证指数', '深证成指', '创业板指', '上证50', '沪深300', '中证500', '中证1000']
index.sort_values(by='ths_pe_ttm_quantile_index', ascending=False, inplace=True)
first_name = index.iloc[0, :]['thscode']
last_name = index.iloc[-1, :]['thscode']
i = index.iloc[0, :]['ths_pe_ttm_quantile_index']
if i > 0 and i <= 0.1:
    up.add_run('所有指数近10年PE分位数均处于10%以下；其中，')
elif i > 0.1 and i <= 0.2:
    up.add_run('所有指数近10年PE分位数均处于20%以下；其中，')
elif i > 0.2 and i <= 0.3:
    up.add_run('所有指数近10年PE分位数均处于30%以下；其中，')
elif i > 0.3 and i <= 0.4:
    up.add_run('所有指数近10年PE分位数均处于40%以下；其中，')
elif i > 0.4 and i <= 0.5:
    up.add_run('所有指数近10年PE分位数均处于50%以下；其中，')
elif i > 0.5 and i <= 0.6:
    up.add_run('所有指数近10年PE分位数均处于60%以下；其中，')
else:
    up.add_run('各指数近10年PE分位数分化较大；其中，')
up.add_run(str(first_name))
up.add_run('估值分位偏高，PE分位数接近')
first_score = "%.0f%%" % (round(index.iloc[0, :]['ths_pe_ttm_quantile_index'] / 10) * 10)
up.add_run(str(first_score))
up.add_run('；')
up.add_run(str(last_name))
up.add_run('估值分位偏低，PE分位数低于')
last_score = "%.0f%%" % (math.ceil(index.iloc[-1, :]['ths_pe_ttm_quantile_index']))
up.add_run(str(last_score))
up.add_run('。以下是申万各行业本周情况。')

#################################申万行业###############################################
index = THS_HQ(
    '801050.SL,801730.SL,801950.SL,801150.SL,801970.SL,801170.SL,801960.SL,801130.SL,801770.SL,801740.SL,801780.SL,801030.SL,801080.SL,801140.SL,801110.SL,801230.SL,801760.SL,801890.SL,801720.SL,801040.SL,801790.SL,801160.SL,801710.SL,801180.SL,801200.SL,801120.SL,801750.SL,801980.SL,801880.SL,801210.SL,801010.SL',
    'changeRatio', 'Interval:W', Mon, Fri).data
index['thscode'] = ['有色金属', '电力设备', '煤炭', '医药生物', '环保', '交通运输', '石油石化', '纺织服饰', '通信',
                    '国防军工', '银行', '基础化工', '电子', '轻工制造', '家用电器', '综合', '传媒', '机械设备',
                    '建筑装饰', '钢铁', '非银金融', '公用事业', '建筑材料', '房地产', '商贸零售', '食品饮料', '计算机',
                    '美容护理', '汽车', '社会服务', '农林牧渔']
photo()
note()
hangye = text('行业层面，本周')
B = index['changeRatio']
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
if n2 == 0:
    hangye.add_run('所有行业呈现上涨状态。')
elif n1 == 0:
    hangye.add_run('所有行业呈现下跌状态。')
elif n1 > n2:
    hangye.add_run('多数行业呈现上涨状态。')
elif n2 > n1:
    hangye.add_run('多数行业呈现下跌状态。')
else:
    hangye.add_run('行业涨跌不一')
hangye.add_run('其中，')
index.sort_values(by='changeRatio', ascending=False, inplace=True)
first_name = index.iloc[0, :]['thscode']
first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
last_name = index.iloc[-1, :]['thscode']
last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
hangye.add_run(str(first_name))
hangye.add_run('行业领涨，周涨')
hangye.add_run(str(first_score))
hangye.add_run('；')
hangye.add_run(str(last_name))
hangye.add_run('行业领跌，周跌')
hangye.add_run(str(last_score))
hangye.add_run('。年度来看，')
index = THS_BD(
    '801050.SL,801730.SL,801950.SL,801150.SL,801970.SL,801170.SL,801960.SL,801130.SL,801770.SL,801740.SL,801780.SL,801030.SL,801080.SL,801140.SL,801110.SL,801230.SL,801760.SL,801890.SL,801720.SL,801040.SL,801790.SL,801160.SL,801710.SL,801180.SL,801200.SL,801120.SL,801750.SL,801980.SL,801880.SL,801210.SL,801010.SL',
    'ths_int_chg_ratio_index', Year).data
index['thscode'] = ['有色金属', '电力设备', '煤炭', '医药生物', '环保', '交通运输', '石油石化', '纺织服饰', '通信',
                    '国防军工', '银行', '基础化工', '电子', '轻工制造', '家用电器', '综合', '传媒', '机械设备',
                    '建筑装饰', '钢铁', '非银金融', '公用事业', '建筑材料', '房地产', '商贸零售', '食品饮料', '计算机',
                    '美容护理', '汽车', '社会服务', '农林牧渔']
index.sort_values(by='ths_int_chg_ratio_index', ascending=False, inplace=True)
first_name = index.iloc[0, :]['thscode']
last_name = index.iloc[-1, :]['thscode']
hangye.add_run(str(first_name))
hangye.add_run('行业领涨，')
hangye.add_run(str(last_name))
hangye.add_run('行业领跌。')
hangye.add_run('从行业估值分位来看，各行业分位分化较大')
index = THS_BD(
    '801050.SL,801730.SL,801950.SL,801150.SL,801970.SL,801170.SL,801960.SL,801130.SL,801770.SL,801740.SL,801780.SL,801030.SL,801080.SL,801140.SL,801110.SL,801230.SL,801760.SL,801890.SL,801720.SL,801040.SL,801790.SL,801160.SL,801710.SL,801180.SL,801200.SL,801120.SL,801750.SL,801980.SL,801880.SL,801210.SL,801010.SL',
    'ths_pe_ttm_quantile_index', Value).data
index['thscode'] = ['有色金属', '电力设备', '煤炭', '医药生物', '环保', '交通运输', '石油石化', '纺织服饰', '通信',
                    '国防军工', '银行', '基础化工', '电子', '轻工制造', '家用电器', '综合', '传媒', '机械设备',
                    '建筑装饰', '钢铁', '非银金融', '公用事业', '建筑材料', '房地产', '商贸零售', '食品饮料', '计算机',
                    '美容护理', '汽车', '社会服务', '农林牧渔']
B = index['ths_pe_ttm_quantile_index']
B_gt_0 = B < 0.5
n1 = B_gt_0.sum()
n = n1 / 31
index.sort_values(by='ths_pe_ttm_quantile_index', ascending=False, inplace=True)
A = index.iloc[0, :]['thscode']
B = index.iloc[1, :]['thscode']
C = index.iloc[2, :]['thscode']
X = index.iloc[-1, :]['thscode']
Y = index.iloc[-2, :]['thscode']
Z = index.iloc[-3, :]['thscode']
if n > 0.5:
    hangye.add_run('，超半数的行业估值分位处于50%以下。其中，估值最高的三个行业为')
else:
    hangye.add_run('。其中，估值最高的三个行业为')
hangye.add_run(str(A))
hangye.add_run('、')
hangye.add_run(str(B))
hangye.add_run('和')
hangye.add_run(str(C))
hangye.add_run('，估值最低的三个行业为')
hangye.add_run(str(X))
hangye.add_run('、')
hangye.add_run(str(Y))
hangye.add_run('和')
hangye.add_run(str(Z))
hangye.add_run('。')
head3('指数成交额')
photo()
note()
a = text('本周市场活跃度有所')
a.add_run('下降')
a.add_run('，沪深两市日均成交额为')
a.add_run('9309.23')
a.add_run('亿，较上周沪深两市日均成交额')
a.add_run('11395.39')
a.add_run('亿有所')
a.add_run('下降')
a.add_run('。市场交易集中度如下图。')
run1 = a.runs[1]
run1.font.color.rgb = RGBColor(255, 0, 0)
run2 = a.runs[3]
run2.font.color.rgb = RGBColor(255, 0, 0)
run3 = a.runs[5]
run3.font.color.rgb = RGBColor(255, 0, 0)
run4 = a.runs[7]
run4.font.color.rgb = RGBColor(255, 0, 0)
head3('交易集中度')
photo()
note()
text('以下3张图统计了全市场超越三大指数涨跌幅的股票数量，以刻画市场超额收益环境。')
head3('超额股票数量占比（沪深300）')
photo()
head3('超额股票数量占比（中证500）')
photo()
head3('超额股票数量占比（中证1000）')
photo()
head3('常用指数换手率')
photo()
head3('指数截面波动率')
photo()
head3('滚动5日波动率')
photo()
head3('滚动10日波动率')
photo()
head3('滚动20日波动率')
photo()
note2()
head3('年化基差')
photo()
note()
a = text('基差方面，本周四大期指基差')
a.add_run('走势基本相似，总体呈现出先下降后上升的趋势。本周各合约基差基本走平，对中性产品的收益无较大影响。')
run1 = a.runs[1]
run1.font.color.rgb = RGBColor(255, 0, 0)
head3('Barra因子走势')
photo()
note()
a = text('风格因子表现方面，本周')
a.add_run('动量因子和beta因子上涨，市值因子、流动性因子和价值因子下跌，其余因子表现较为平淡。')
run1 = a.runs[1]
run1.font.color.rgb = RGBColor(255, 0, 0)
#################################债券市场###############################################
head2('（二）债券市场')
photo()
note()
index = THS_HQ('h11001.CSI,h11006.CSI,h11073.CSI,000832.CSI', 'changeRatio', 'Interval:W', Mon, Fri).data
index['thscode'] = ['中证全债', '中证国债', '中证信用', '中证转债']
B = index['changeRatio']
zhai = text('债券方面，本周各债券指数')
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
index.sort_values(by='changeRatio', ascending=False, inplace=True)
first_name = index.iloc[0, :]['thscode']
first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
last_name = index.iloc[-1, :]['thscode']
last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
if n2 == 1:  # 只有一个跌
    # 降序,筛选极值
    zhai.add_run('多数上涨。其中，')
    zhai.add_run(str(first_name))
    zhai.add_run('领涨，周涨')
    zhai.add_run(str(first_score))
    zhai.add_run('；仅')
    zhai.add_run(str(last_name))
    zhai.add_run('下跌，周跌')
    zhai.add_run(str(last_score))
elif n1 == 1:  # 只有一个涨
    # 升序,筛选极值
    zhai.add_run('多数下跌。其中，')
    zhai.add_run(str(last_name))
    zhai.add_run('领跌，周跌')
    zhai.add_run(str(last_score))
    zhai.add_run('；仅')
    zhai.add_run(str(first_name))
    zhai.add_run('上涨，周涨')
    zhai.add_run(str(first_score))
elif n2 == 0:  # 没有跌
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * 100)
    zhai.add_run('均有所上涨。其中，')
    zhai.add_run(str(first_name))
    zhai.add_run('领涨，周涨')
    zhai.add_run(str(first_score))
    zhai.add_run('；')
    zhai.add_run(str(last_name))
    zhai.add_run('涨幅最小，周涨')
    zhai.add_run(str(last_score))
elif n1 == 0:  # 没有涨
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'] * -100)
    zhai.add_run('均有所下跌。其中，')
    zhai.add_run(str(first_name))
    zhai.add_run('领跌，周跌')
    zhai.add_run(str(first_score))
    zhai.add_run('；')
    zhai.add_run(str(last_name))
    zhai.add_run('跌幅最小，周跌')
    zhai.add_run(str(last_score))
else:
    zhai.add_run('涨跌不一。周度来看，')
    zhai.add_run(str(first_name))
    zhai.add_run('领涨，周涨')
    zhai.add_run(str(first_score))
    zhai.add_run('；')
    zhai.add_run(str(last_name))
    zhai.add_run('领跌，周跌')
    zhai.add_run(str(last_score))
zhai.add_run('。')
photo()
note()
a = text(
    '利率角度，本周资金面整体维持宽松，DR007周内先下后上，SHIBOR 3M利率周内回落4BP来到2.24%附近，市场利率整体走平。近期利率债震荡走平，10年期国债利率在2.70-2.71%区间内窄幅波动，流动性宽松带动短端利率下行，1年期国债利率降至2.0%以下。')
run1 = a.runs[0]
run1.font.color.rgb = RGBColor(255, 0, 0)
#################################商品市场###############################################
head2('（三）商品市场')
photo()
note()
head3('大宗商品走势')
photo()
note()
index = THS_HQ('nhci.sl,nhai.sl,nhnfi.sl,nheci.sl,nhpmi.sl,nhfi.sl', 'changeRatio', 'Interval:W', Mon, Fri).data
index['thscode'] = ['南华商品', '南华农产品', '南华有色金属', '南华能化', '南华贵金属', '南华黑色']
B = index['changeRatio']
shang = text('本周商品市场各品种')
nh = index.iloc[0, :]['changeRatio']
nanhua = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
nanhua1 = "%.2f%%" % (index.iloc[0, :]['changeRatio'] * -1)
index.drop(index=0)
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
if n2 == 0:  # 没有跌
    index.sort_values(by='changeRatio', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
    second_name = index.iloc[1, :]['thscode']
    second_score = "%.2f%%" % (index.iloc[1, :]['changeRatio'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'])
    shang.add_run('全部上涨，南华商品指数周涨')
    shang.add_run(str(nanhua))
    shang.add_run('。其中，')
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，周涨')
    shang.add_run(str(first_score))
    shang.add_run('；其次是')
    shang.add_run(str(second_name))
    shang.add_run('指数，周涨')
    shang.add_run(str(second_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数涨幅最小，周涨')
    shang.add_run(str(last_score))
elif n1 == 0:  # 没有涨
    # 降序,筛选极值
    index.sort_values(by='changeRatio', ascending=True, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'] * -1)
    second_name = index.iloc[1, :]['thscode']
    second_score = "%.2f%%" % (index.iloc[1, :]['changeRatio'] * -1)
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    shang.add_run('全部下跌，南华商品指数周跌')
    shang.add_run(str(nanhua1))
    shang.add_run('。其中，')
    shang.add_run(str(first_name))
    shang.add_run('指数领跌，周跌')
    shang.add_run(str(first_score))
    shang.add_run('；其次是')
    shang.add_run(str(second_name))
    shang.add_run('指数，周跌')
    shang.add_run(str(second_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数跌幅最小，周跌')
    shang.add_run(str(last_score))
elif n2 == 1:  # 只有一个跌
    # 降序,筛选极值
    index.sort_values(by='changeRatio', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    shang.add_run('多数上涨，南华商品指数周涨')
    shang.add_run(str(nanhua))
    shang.add_run('。其中，')
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，周涨')
    shang.add_run(str(first_score))
    shang.add_run('；仅')
    shang.add_run(str(last_name))
    shang.add_run('指数下跌，周跌')
    shang.add_run(str(last_score))
elif n1 == 1:  # 只有一个涨
    # 升序,筛选极值
    index.sort_values(by='changeRatio', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    shang.add_run('多数下跌，南华商品指数周跌')
    shang.add_run(str(nanhua1))
    shang.add_run('。其中，')
    shang.add_run(str(last_name))
    shang.add_run('指数领跌，周跌')
    shang.add_run(str(last_score))
    shang.add_run('；仅')
    shang.add_run(str(first_name))
    shang.add_run('指数上涨，周涨')
    shang.add_run(str(first_score))
else:
    index.sort_values(by='changeRatio', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['changeRatio'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['changeRatio'] * -1)
    shang.add_run('涨跌不一，南华商品指数')
    if nh > 0:
        shang.add_run('上涨')
        shang.add_run(str(nanhua))
    elif nh < 0:
        shang.add_run('下跌')
        shang.add_run(str(nanhua1))
    else:
        shang.add_run('本周走平')
    shang.add_run('。其中，')
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，周涨')
    shang.add_run(str(first_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数领跌，周跌')
    shang.add_run(str(last_score))
shang.add_run('。年度来看，')
index = THS_BD('nhci.sl,nhai.sl,nhnfi.sl,nheci.sl,nhpmi.sl,nhfi.sl', 'ths_int_chg_ratio_index', Year).data
index['thscode'] = ['南华商品', '南华农产品', '南华有色金属', '南华能化', '南华贵金属', '南华黑色']
B = index['ths_int_chg_ratio_index']
B_gt_0 = B > 0
n1 = B_gt_0.sum()
B_gt_0 = B < 0
n2 = B_gt_0.sum()
if n2 == 0:  # 没有跌
    index.sort_values(by='ths_int_chg_ratio_index', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['ths_int_chg_ratio_index'])
    second_name = index.iloc[1, :]['thscode']
    second_score = "%.2f%%" % (index.iloc[1, :]['ths_int_chg_ratio_index'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['ths_int_chg_ratio_index'])
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，年初至今上涨')
    shang.add_run(str(first_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数涨幅最小，年初至今上涨')
    shang.add_run(str(last_score))
elif n1 == 0:  # 没有涨
    # 降序,筛选极值
    index.sort_values(by='ths_int_chg_ratio_index', ascending=True, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['ths_int_chg_ratio_index'] * -1)
    second_name = index.iloc[1, :]['thscode']
    second_score = "%.2f%%" % (index.iloc[1, :]['ths_int_chg_ratio_index'] * -1)
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['ths_int_chg_ratio_index'] * -1)
    shang.add_run(str(first_name))
    shang.add_run('指数领跌，年初至今下跌')
    shang.add_run(str(first_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数跌幅最小，年初至今下跌')
    shang.add_run(str(last_score))
elif n2 == 1:  # 只有一个跌
    # 降序,筛选极值
    index.sort_values(by='ths_int_chg_ratio_index', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['ths_int_chg_ratio_index'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['ths_int_chg_ratio_index'] * -1)
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，年初至今上涨')
    shang.add_run(str(first_score))
    shang.add_run('；仅')
    shang.add_run(str(last_name))
    shang.add_run('指数下跌，年初至今下跌')
    shang.add_run(str(last_score))
elif n1 == 1:  # 只有一个涨
    # 升序,筛选极值
    index.sort_values(by='ths_int_chg_ratio_index', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['ths_int_chg_ratio_index'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['ths_int_chg_ratio_index'] * -1)
    shang.add_run(str(last_name))
    shang.add_run('指数领跌，年初至今下跌')
    shang.add_run(str(last_score))
    shang.add_run('；仅')
    shang.add_run(str(first_name))
    shang.add_run('指数上涨，年初至今上涨')
    shang.add_run(str(first_score))
else:
    index.sort_values(by='ths_int_chg_ratio_index', ascending=False, inplace=True)
    first_name = index.iloc[0, :]['thscode']
    first_score = "%.2f%%" % (index.iloc[0, :]['ths_int_chg_ratio_index'])
    last_name = index.iloc[-1, :]['thscode']
    last_score = "%.2f%%" % (index.iloc[-1, :]['ths_int_chg_ratio_index'] * -1)
    shang.add_run(str(first_name))
    shang.add_run('指数领涨，年初至今上涨')
    shang.add_run(str(first_score))
    shang.add_run('；')
    shang.add_run(str(last_name))
    shang.add_run('指数领跌，年初至今下跌')
    shang.add_run(str(last_score))
shang.add_run('。')

#####################################第二部分####################################################
head1('二、市场及策略观察')
head2('（一）市场周度观察')
a = text('本周（')
a.add_run(str(monday)).bold = True
a.add_run('-').bold = True
a.add_run(str(friday)).bold = True
a.add_run('）')
a = text('本周（')
a.add_run(str(monday)).bold = True
a.add_run('-').bold = True
a.add_run(str(friday)).bold = True
a.add_run('）')
a = text('本周（')
a.add_run(str(monday)).bold = True
a.add_run('-').bold = True
a.add_run(str(friday)).bold = True
a.add_run('）')

document.add_section(WD_SECTION.NEW_PAGE)
a = head3('免责声明')
a.alignment = WD_ALIGN_PARAGRAPH.LEFT
a.paragraph_format.space_after = Pt(8)
a.paragraph_format.space_before = Pt(0)
text('本内容的观点和信息仅供内部阅读，请勿对外公开传播和发放。')
text(
    '本内容不构成期货投资咨询服务，不构成具体业务或产品的推介，亦不应被视为相应金融产品的投资建议。市场有风险，投资需谨慎。投资者不应将本内容作为做出投资决策的参考因素，亦不应认为本内容可以取代自己的判断。在决定投资前，如有需要，投资者务必向专业人士咨询并谨慎决策。')
text(
    '本内容的信息来源于公开资料及已获授权的资料，本公司对该等信息的准确性、完整性或可靠性不作任何保证。本内容的撰写力求独立、客观和公正，结论不受任何第三方的授意或影响。本内容所载的资料、意见及推测仅反映本公司于发布本内容当日的判断，本内容所指的产品的价格可升可跌，过往表现不应作为日后的表现依据。在不同时期，本公司可发出与本内容所载资料、意见及推测不一致的内容。本公司不保证本内容所含信息保持在最新状态。同时，本公司对本内容所含信息可在不发出通知的情形下做出修改，投资者应当自行关注相应的更新或修改。')
text(
    '本内容版权仅为本公司所有，未经书面许可，任何机构和个人不得以任何形式翻版、复制、转发和发布，且不得对本内容进行有悖原意的引用、删节和修改。')

for paragraph in document.paragraphs:
    for run in paragraph.runs:
        run.font.name = '仿宋'
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), '仿宋')
p = document.paragraphs[0]
run = p.runs[1]
run.font.name = '方正兰亭粗黑简体'
r = run._element.rPr.rFonts
r.set(qn('w:eastAsia'), '方正兰亭粗黑简体')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(254, 230, 206)
# 输出word
name = 'output/天风50私享-市场观察周报（' + Monday + '-' + Friday + '）' + '.docx'
document.save(name)
